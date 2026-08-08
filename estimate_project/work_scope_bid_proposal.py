from datetime import datetime
from tkinter import messagebox, Text, Checkbutton, BooleanVar
from tkcalendar import DateEntry
import tkinter as tk
import os
import logging
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt
import mplcursors
from tkinter import *
from tkinter import ttk, messagebox
from sklearn.model_selection import train_test_split, GridSearchCV, RandomizedSearchCV, KFold
from sklearn.ensemble import RandomForestRegressor
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.linear_model import Ridge
from datetime import datetime
from idlelib.tooltip import Hovertip

from building_demo import BuildingDemo
from interior_demo import InteriorDemo
from house_demo import HouseDemo
from contact_book import ContactBook
from equipment_book import EquipmentBook
from postgresql import get_db_connection


logger = logging.getLogger(__name__)

# ============================================================================ #
# ================================= INFO ===================================== #
# ============================================================================ #

# ============================================================================ #
# ================================= TODO ===================================== #
# ============================================================================ #
# TODO: 
# Add PostgreSQL variables to .env before running program. 
# See postgresql.py for required parameters.
# ============================================================================ #


FONT = "Times New Roman"
FONT_SIZE = 13
TITLE_FONT = "Arial"
TITLE_FONT_SIZE = 16
PAGE_HEIGHT = 35
PAGE_WIDTH = 75

collected_data = {}

todays_date = datetime.now().strftime("%m%d%Y")
downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")

# =========================================================================== #
# ======================== Get Data from PostgreSQL ========================= #
# =========================================================================== #

def fetch_data_from_postgresql():
    """Fetch project data from PostgreSQL and shape it for model training."""
    query = """
        SELECT
            job_number AS "Job Number",
            awarded_date AS "Awarded Date",
            project_description AS "Description",
            structure_type AS "Structure Type",
            sqft AS "SqFt",
            bid_price AS "Bid Price",
            job_cost AS "Job Cost",
            estimator AS "Estimator"
        FROM project;
    """

    with get_db_connection() as conn:
        df = pd.read_sql_query(query, conn)

    if df.empty:
        messagebox.showwarning(
            "No Project Data",
            "No project rows were returned from PostgreSQL table 'project'.",
        )
        return df

    df['SqFt'] = pd.to_numeric(df['SqFt'], errors='coerce')
    df['Bid Price'] = pd.to_numeric(df['Bid Price'], errors='coerce')
    df['Job Cost'] = pd.to_numeric(df['Job Cost'], errors='coerce')
    df['Awarded Date'] = pd.to_datetime(df['Awarded Date'], errors='coerce')

    # Drop rows with missing model-critical values to keep downstream training stable.
    df = df.dropna(subset=['Job Number', 'Description', 'Structure Type', 'SqFt', 'Bid Price', 'Job Cost'])
    df['SqFt'] = df['SqFt'].astype(int)
    df['Profit and Loss %'] = round(((df['Bid Price'] - df['Job Cost']) / df['Bid Price']) * 100, 2)
    df = df.set_index('Job Number')

    print("\nData from PostgreSQL")
    print("\nDataframe with 'Profit and Loss %' column:")
    print(df.tail())
    return df

# =========================================================================== #
# ======================= New Estimate costs using models =================== #
# =========================================================================== #

def train_models(df):
    """
    Train two tuned ML model families for bid and cost prediction.

    Models:
    - Ridge Regression: regularized linear baseline.
    - Random Forest: non-linear ensemble with constraints to reduce overfitting.

    Both models use cross-validation and hyperparameter tuning.

    Parameters:
    df (DataFrame): The input DataFrame containing the features and target variables.

        Returns:
        tuple: Tuned estimators, holdout sets, and CV summary values.
    """
    X = df[['Description', 'Structure Type', 'SqFt']]
    y_bid = df['Bid Price']
    y_cost = df['Job Cost']

    # Split the data into training and testing sets
    X_train, X_test, y_bid_train, y_bid_test, y_cost_train, y_cost_test = \
        train_test_split(X, y_bid, y_cost, test_size=0.2, random_state=42)
    
    # Define the column transformer for one-hot encoding
    preprocessor = ColumnTransformer(
        transformers=[
            ('cat', OneHotEncoder(), ['Description', 'Structure Type']),
            ('num', StandardScaler(), ['SqFt'])
        ])

    # Use a dynamic CV split count so this also works with smaller datasets.
    cv_splits = min(5, max(2, len(X_train) // 4))
    cv_strategy = KFold(n_splits=cv_splits, shuffle=True, random_state=42)

    # Ridge (regularized linear model) with CV tuning.
    ridge_bid_pipeline = Pipeline(steps=[
        ('preprocessor', preprocessor),
        ('model', Ridge())
    ])
    ridge_cost_pipeline = Pipeline(steps=[
        ('preprocessor', preprocessor),
        ('model', Ridge())
    ])
    ridge_params = {
        'model__alpha': [0.01, 0.1, 1.0, 5.0, 10.0, 25.0, 50.0, 100.0]
    }

    ridge_bid_search = GridSearchCV(
        ridge_bid_pipeline,
        ridge_params,
        cv=cv_strategy,
        scoring='neg_mean_absolute_error',
        n_jobs=1,
    )
    ridge_bid_search.fit(X_train, y_bid_train)

    ridge_cost_search = GridSearchCV(
        ridge_cost_pipeline,
        ridge_params,
        cv=cv_strategy,
        scoring='neg_mean_absolute_error',
        n_jobs=1,
    )
    ridge_cost_search.fit(X_train, y_cost_train)

    # Random Forest with conservative defaults and CV tuning to reduce overfitting.
    rf_bid_pipeline = Pipeline(steps=[
        ('preprocessor', preprocessor),
        ('model', RandomForestRegressor(random_state=42, n_jobs=1))
    ])
    rf_cost_pipeline = Pipeline(steps=[
        ('preprocessor', preprocessor),
        ('model', RandomForestRegressor(random_state=42, n_jobs=1))
    ])
    rf_params = {
        'model__n_estimators': [200, 350, 500],
        'model__max_depth': [5, 8, 12, None],
        'model__min_samples_split': [2, 4, 6, 10],
        'model__min_samples_leaf': [1, 2, 4],
        'model__max_features': ['sqrt', 0.7, 1.0],
        'model__bootstrap': [True]
    }

    rf_bid_search = RandomizedSearchCV(
        rf_bid_pipeline,
        rf_params,
        n_iter=15,
        cv=cv_strategy,
        scoring='neg_mean_absolute_error',
        n_jobs=1,
        random_state=42,
    )
    rf_bid_search.fit(X_train, y_bid_train)

    rf_cost_search = RandomizedSearchCV(
        rf_cost_pipeline,
        rf_params,
        n_iter=15,
        cv=cv_strategy,
        scoring='neg_mean_absolute_error',
        n_jobs=1,
        random_state=42,
    )
    rf_cost_search.fit(X_train, y_cost_train)

    cv_results = {
        'cv_splits': cv_splits,
        'ridge_bid_cv_mae': abs(ridge_bid_search.best_score_),
        'ridge_cost_cv_mae': abs(ridge_cost_search.best_score_),
        'rf_bid_cv_mae': abs(rf_bid_search.best_score_),
        'rf_cost_cv_mae': abs(rf_cost_search.best_score_),
        'ridge_bid_best_params': ridge_bid_search.best_params_,
        'ridge_cost_best_params': ridge_cost_search.best_params_,
        'rf_bid_best_params': rf_bid_search.best_params_,
        'rf_cost_best_params': rf_cost_search.best_params_,
    }

    return ridge_bid_search.best_estimator_, ridge_cost_search.best_estimator_, \
        rf_bid_search.best_estimator_, rf_cost_search.best_estimator_, X_test, \
        y_bid_test, y_cost_test, cv_results

# ===========================================================================
# ======================= New Estimate costs using models ===================
# ===========================================================================
def estimate_costs(models, square_feet, description, structure_type):
    """
    Estimate bid prices and job costs using multiple machine learning models.

    Parameters:
    models (list): A list containing trained models in this order:
    [ridge_bid_model, ridge_cost_model, rf_bid_model, rf_cost_model]
    square_feet (float): The square footage of the project.
    description (str): A description of the project(e.g., Building Demo, 
    House Demo, Interior Demolition).
    structure_type (str): The type of structure (e.g., Wood, Metal, Other).

    Returns:
    tuple: A tuple containing four estimated values in this order:
        (ridge_estimated_bid_price, ridge_estimated_job_cost,
            rf_estimated_bid_price, rf_estimated_job_cost)
    """
    ridge_bid_model, ridge_cost_model, rf_bid_model, rf_cost_model = models[:4]
    input_data = pd.DataFrame([[description, structure_type, square_feet]], 
                            columns=['Description', 'Structure Type', 'SqFt'])
    ridge_estimated_bid_price = ridge_bid_model.predict(input_data)[0]
    ridge_estimated_job_cost = ridge_cost_model.predict(input_data)[0]
    rf_estimated_bid_price = rf_bid_model.predict(input_data)[0]
    rf_estimated_job_cost = rf_cost_model.predict(input_data)[0]

    return ridge_estimated_bid_price, ridge_estimated_job_cost, \
        rf_estimated_bid_price, rf_estimated_job_cost

# ===========================================================================
# ========================== Main Program =================================== 
# ===========================================================================
def on_closing_estimates(info):
    if info:
        info.quit()
        info.destroy()


def create_tooltip(widget, text):
    """Creates tooltips for the tkinter launch_demo_window widgets"""
    Hovertip(widget, text, hover_delay=500)


def save_to_word(data, file_path):
    """
    Save the collected data to a Word document.

    Parameters:
    data (dict): A dictionary containing the collected data.
    file_path (str): The path to save the Word document.

    Returns:
    None
    """
    document = Document()
    document.add_heading("Work Scope Bid Proposal", level=1)

    for section, content in data.items():
        # Add a heading for each section
        document.add_heading(section, level=2)

        # Add the starting text as a paragraph
        if "starting_text" in content:
            document.add_paragraph(content["starting_text"])

        # Add the user input as a paragraph
        if "user_input" in content:
            document.add_paragraph(content["user_input"])

        # Add the total cost if it exists
        if "total_cost" in content:
            document.add_paragraph(content["total_cost"])

    # Save the document to the specified file path
    document.save(file_path)
    messagebox.showinfo("Success", f"Proposal saved to {file_path}")


def create_navigation_menu(root):
    """
    Create a navigation menu for the application.

    This function adds a menu bar to the root window, allowing users to navigate
    directly to any page.

    Parameters:
    root (Tk): The root Tkinter window.

    Returns:
    None
    """
    menu_bar = tk.Menu(root)

    # Create a "Navigate" menu
    navigate_menu = tk.Menu(menu_bar, tearoff=0)
    navigate_menu.add_command(label="Project Overview", 
                              command=lambda: show_page("Project Overview"))
    navigate_menu.add_command(label="Estimating Details",
                              command=lambda: show_page("Estimating Details"))
    navigate_menu.add_command(label="Project-Specific Details", 
                              command=lambda: show_page("Project-Specific Details"))
    navigate_menu.add_command(label="Methodology and Approach", 
                              command=lambda: show_page("Methodology and Approach"))
    navigate_menu.add_command(label="Order of Operations", 
                              command=lambda: show_page("Order of Operations"))
    navigate_menu.add_command(label="Alternates and Breakouts", 
                              command=lambda: show_page("Alternates and Breakouts"))
    navigate_menu.add_command(label="Inclusions", 
                              command=lambda: show_page("Inclusions"))
    navigate_menu.add_command(label="Exclusions", 
                              command=lambda: show_page("Exclusions"))
    navigate_menu.add_command(label="Equipment", 
                              command=lambda: show_page("Equipment"))
    navigate_menu.add_command(label="Terms and Conditions", 
                              command=lambda: show_page("Terms and Conditions"))
    navigate_menu.add_command(label="Regulatory Compliance", 
                              command=lambda: show_page("Regulatory Compliance"))
    navigate_menu.add_command(label="Final Checklist", 
                              command=lambda: show_page("Final Checklist"))
    navigate_menu.add_command(label="Final Page", 
                              command=lambda: show_page("Final Page"))

    # Add the "Navigate" menu to the menu bar
    menu_bar.add_cascade(label="Navigate", menu=navigate_menu)

    # Configure the menu bar in the root window
    root.config(menu=menu_bar)


def show_toast(message, message_type="info"):
    """
    Show a toast-like notification.

    Parameters:
    message (str): The message to display.
    message_type (str): The type of message ("info", "warning", "error").

    Returns:
    None
    """
    toast = tk.Toplevel(root)
    toast.overrideredirect(True)  # Remove window decorations
    toast.geometry("300x50+500+300")  # Set size and position
    toast.attributes("-topmost", True)  # Keep on top

    # Set background color based on message type
    bg_color = "green" if message_type == "info" else "orange" if message_type == "warning" else "red"

    tk.Label(toast, text=message, bg=bg_color, fg="black", font=(FONT, FONT_SIZE)).pack(fill="both", expand=True)

    # Auto-dismiss the toast after 3 seconds
    # toast.after(3000, toast.destroy)

    # Auto-dismiss the toast after 3 seconds, but check if it still exists
    def safe_destroy():
        try:
            toast.destroy()
        except tk.TclError:
            pass

    toast.after(3000, safe_destroy)


contact_book = ContactBook()

# ============================================================================ #
# ============================= Main GUI ===================================== #
# ============================================================================ #
root = tk.Tk()
root.title("Work Scope Bid Proposal")
root.minsize(800, 800)
root.config(padx=125, pady=50)

# Add the navigation menu
create_navigation_menu(root)

# Create a status bar
# status_label = tk.Label(root, text="", bd=1, relief="sunken", anchor="w", font=(FONT, FONT_SIZE))
# status_label.grid(row=1, column=0, columnspan=2, sticky="ew", padx=10, pady=5)

# Dictionary to store user input
collected_data = {}

# Create a container for all pages
pages = {}

# Function to switch between pages
def show_page(page_name):
    """
    Select a contact from the contact list and populate the main GUI fields.

    This function retrieves the selected contact from the contact list and 
    populates the corresponding fields in the main GUI (e.g., contact name, 
    company, phone, and email). If the contact list is filtered, it uses the 
    filtered data; otherwise, it loads the complete contact data. If no 
    contact is selected, it displays a warning message.

    Parameters:
    None

    Returns:
    None
    """
    for page in pages.values():
        page.grid_forget()  # Hide all pages
    pages[page_name].grid(row=0, column=0, sticky="nsew")  # Show the selected page


# Page 1: Project Overview
def create_project_overview_page():
    """
    Create the "Project Overview" page in the main GUI.

    This function creates the "Project Overview" page, which allows the user 
    to input details such as the project name, address, bid due date, project 
    type, and contact information. The data entered on this page is saved to 
    the `collected_data` dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    global contact_name_var, company_var, phone_var, email_var, project_name_var
    global billing_address_var

    page = tk.Frame(root, padx=75, pady=75)
    pages["Project Overview"] = page

    tk.Label(page, text="Project Overview", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Input fields for project overview
    tk.Label(page, text="Bid Form (Yes/No):").grid(row=1, column=0, sticky="e", padx=10, pady=5)
    bid_form_var = tk.StringVar(value="Yes")
    tk.OptionMenu(page, bid_form_var, "Yes", "No").grid(row=1, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Grading Permit:").grid(row=2, column=0, sticky="e", padx=10, pady=5)
    grading_permit_var = tk.StringVar(value="Included")
    tk.OptionMenu(page, grading_permit_var, "Included", "Excluded", "N/A").grid(row=2, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Project Name:").grid(row=3, column=0, sticky="e", padx=10, pady=5)
    project_name_var = tk.StringVar()
    tk.Entry(page, textvariable=project_name_var, width=40).grid(row=3, column=1, padx=10, pady=5)

    tk.Label(page, text="Project Address:").grid(row=4, column=0, sticky="e", padx=10, pady=5)
    project_address_var = tk.StringVar()
    tk.Entry(page, textvariable=project_address_var, width=40).grid(row=4, column=1, padx=10, pady=5)

    tk.Label(page, text="Bid Due:").grid(row=5, column=0, sticky="e", padx=10, pady=5)
    bid_due_var = tk.StringVar()
    DateEntry(page, textvariable=bid_due_var, date_pattern="mm/dd/yyyy").grid(row=5, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Project Type:").grid(row=6, column=0, sticky="e", padx=10, pady=5)
    project_type_var = tk.StringVar(value="Residential Interior Remodel")
    tk.OptionMenu(page, project_type_var, 
                  "Residential Interior Remodel", 
                  "Residential Complete Demolition (Wrecking)", 
                  "Commercial Interior TI", 
                  "Commercial Complete Demolition (Incl. Slabs/Foundations)", 
                  "Commercial Demo (Top of SOG Only)").grid(row=6, column=1, sticky="w", padx=10, pady=5)
        
    tk.Button(
    page,
    text="Open Contact Book",
    command=lambda: contact_book.open_contact_book(
        company_var,
        billing_address_var,
        contact_name_var,
        phone_var,
        email_var
    )
    ).grid(row=7, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Company:").grid(row=8, column=0, sticky="e", padx=10, pady=5)
    company_var = tk.StringVar()
    tk.Entry(page, textvariable=company_var, width=40).grid(row=8, column=1, padx=10, pady=5)

    tk.Label(page, text="Billing Address:").grid(row=9, column=0, sticky="e", padx=10, pady=5)
    billing_address_var = tk.StringVar()
    tk.Entry(page, textvariable=billing_address_var, width=40).grid(row=9, column=1, padx=10, pady=5)    

    tk.Label(page, text="Contact Name:").grid(row=10, column=0, sticky="e", padx=10, pady=5)
    contact_name_var = tk.StringVar()
    tk.Entry(page, textvariable=contact_name_var, width=40).grid(row=10, column=1, padx=10, pady=5)

    tk.Label(page, text="Phone:").grid(row=11, column=0, sticky="e", padx=10, pady=5)
    phone_var = tk.StringVar()
    tk.Entry(page, textvariable=phone_var, width=40).grid(row=11, column=1, padx=10, pady=5)

    tk.Label(page, text="Email:").grid(row=12, column=0, sticky="e", padx=10, pady=5)
    email_var = tk.StringVar()
    tk.Entry(page, textvariable=email_var, width=40).grid(row=12, column=1, padx=10, pady=5)

    def save_project_overview():
        """
        Save the data entered on the "Project Overview" page.

        This function collects the user input from the "Project Overview" page, 
        formats it, and saves it to the `collected_data` dictionary under the 
        "Project Overview" key. After saving, it navigates to the 
        "Project-Specific Details" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect all input values
        bid_form_value = bid_form_var.get()
        grading_permit_value = grading_permit_var.get()
        project_name_value = project_name_var.get()
        project_address_value = project_address_var.get()
        bid_due_value = bid_due_var.get()
        project_type_value = project_type_var.get()
        company_value = company_var.get()  
        billing_address_value = billing_address_var.get()      
        contact_name_value = contact_name_var.get()
        phone_value = phone_var.get()
        email_value = email_var.get()

        # Format the collected data
        project_overview_data = (
            f"Todays Date: {datetime.now().strftime('%m/%d/%Y')}\n"
            f"Bid Form: {bid_form_value}\n"
            f"Grading Permit: {grading_permit_value}\n"
            f"Project Name: {project_name_value}\n"
            f"Project Address: {project_address_value}\n"
            f"Bid Due: {bid_due_value}\n"
            f"Project Type: {project_type_value}\n"
            f"Company: {company_value}\n"
            f"Billing Address: {billing_address_value}\n"
            f"Contact Name: {contact_name_value}\n"
            f"Phone: {phone_value}\n"
            f"Email: {email_value}\n"
        )

        # Save the data to the collected_data dictionary
        collected_data["Project Overview"] = {
            "user_input": project_overview_data,
        }

        # Proceed to the next page
        show_page("Estimating Details")

    tk.Button(page, text="Next", 
              command=save_project_overview).grid(row=13, column=1, pady=20, sticky="e")


def create_project_estimating_page():
    page = tk.Frame(root)
    pages["Estimating Details"] = page

    tk.Label(page, text="Project-Specific Estimating Details",
            font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)
    tk.Label(
        page,
        text=(
            "Use this page to run historical + ML estimating and save the results "
            "into the proposal document's Estimating Details section."
        ),
        font=(FONT, FONT_SIZE),
        wraplength=700,
        fg="darkgreen",
        justify="left",
    ).grid(row=1, column=0, columnspan=2, padx=10, pady=5, sticky="w")

    model_cache = {"df": None, "models": None}
    estimate_status_var = tk.StringVar(
        value="No estimate saved yet. Complete inputs and click 'Run Estimate'."
    )

    description_var = tk.StringVar()
    structure_var = tk.StringVar()
    sqft_var = tk.StringVar()
    lower_limit_var = tk.StringVar()
    upper_limit_var = tk.StringVar()
    equipment_cost_var = tk.StringVar()
    disposal_cost_var = tk.StringVar()
    num_days_var = tk.StringVar()
    num_workers_var = tk.StringVar()

    def ensure_model_data_loaded():
        if model_cache["df"] is not None and model_cache["models"] is not None:
            return True

        try:
            model_df = fetch_data_from_postgresql()
            if model_df.empty:
                messagebox.showwarning("No Project Data", 
                                       "No project rows were returned from PostgreSQL table 'project'.")
                return False
            model_cache["df"] = model_df
            model_cache["models"] = train_models(model_df)
            return True
        except Exception:
            logger.exception("Unable to load estimating data")
            messagebox.showerror(
                "Estimating Data Error",
                "Unable to load estimating data. Please verify database settings and try again.",
            )
            return False

    def on_description_change(event=None):
        if description_var.get() == "Interior Demolition":
            structure_var.set("Other")

    def _safe_profit_percent(bid_price, job_cost):
        if bid_price == 0:
            return 0.0
        return round(((bid_price - job_cost) / bid_price) * 100, 2)

    def build_project_report(description_value, structure_value, sqft_value, equipment_cost_value,
                             disposal_cost_value, num_days_value, num_workers_value):
        """Build class-based demolition report text for Word output."""
        if description_value == "Interior Demolition":
            project = InteriorDemo(
                description_value,
                structure_value,
                total_sqft=sqft_value,
                total_equipment_cost=equipment_cost_value,
                total_disposal_cost=disposal_cost_value,
                num_days=num_days_value,
                num_guys=num_workers_value,
            )
        elif description_value == "House Demo":
            project = HouseDemo(
                description_value,
                structure_value,
                total_sqft=sqft_value,
                total_equipment_cost=equipment_cost_value,
                total_disposal_cost=disposal_cost_value,
                num_days=num_days_value,
                num_guys=num_workers_value,
            )
        else:
            project = BuildingDemo(
                description_value,
                structure_value,
                total_sqft=sqft_value,
                total_equipment_cost=equipment_cost_value,
                total_disposal_cost=disposal_cost_value,
                num_days=num_days_value,
                num_guys=num_workers_value,
            )

        if not project.validate_input_data():
            return None

        return project.generate_detailed_report()

    def run_estimate_and_store():
        description_value = description_var.get().strip()
        structure_value = structure_var.get().strip()
        equipment_cost = equipment_cost_var.get().strip()
        disposal_cost = disposal_cost_var.get().strip()
        num_days = num_days_var.get().strip()
        num_workers = num_workers_var.get().strip()

        if not description_value:
            messagebox.showerror("Missing Description", "Please select a project description.")
            return

        try:
            sqft_value = int(sqft_var.get().strip())
            lower_limit_value = int(lower_limit_var.get().strip())
            upper_limit_value = int(upper_limit_var.get().strip())
        except ValueError:
            messagebox.showerror("Invalid Input", "Please enter valid numbers for total SqFt and limits.")
            return

        try:
            equipment_cost_value = float(equipment_cost)
            disposal_cost_value = float(disposal_cost)
            num_days_value = int(num_days)
            num_workers_value = int(num_workers)
        except ValueError:
            messagebox.showerror(
                "Invalid Input",
                "Please enter valid numbers for equipment/disposal cost, days, and workers.",
            )
            return

        if sqft_value < 0 or num_days_value <= 0 or num_workers_value <= 0:
            messagebox.showerror(
                "Invalid Input",
                "SqFt, number of days, and number of workers must be greater than zero.",
            )
            return

        if equipment_cost_value < 0 or disposal_cost_value < 0:
            messagebox.showerror(
                "Invalid Input",
                "Equipment and disposal costs must be zero or greater.",
            )
            return

        if upper_limit_value < lower_limit_value:
            messagebox.showerror("Invalid Range", "Upper limit must be greater than or equal to lower limit.")
            return

        if description_value in ["Building Demo", "House Demo"] and not structure_value:
            messagebox.showerror("Missing Structure Type", "Please select a structure type.")
            return

        if not ensure_model_data_loaded():
            return

        model_df = model_cache["df"]
        models = model_cache["models"]

        if description_value in ["Building Demo", "House Demo"]:
            projects = model_df.loc[
                (model_df["Description"] == description_value)
                & (model_df["Structure Type"] == structure_value)
                & (model_df["SqFt"].between(lower_limit_value, upper_limit_value))
            ]
        else:
            projects = model_df.loc[
                (model_df["Description"] == description_value)
                & (model_df["SqFt"].between(lower_limit_value, upper_limit_value))
            ]

        if projects.empty:
            messagebox.showwarning("No Results", "No historical projects matched this estimating range.")
            return

        detailed_project_report = build_project_report(
            description_value,
            structure_value,
            sqft_value,
            equipment_cost_value,
            disposal_cost_value,
            num_days_value,
            num_workers_value,
        )
        if not detailed_project_report:
            messagebox.showerror("Invalid Input", "Unable to generate the detailed project report.")
            return

        average_job_cost = round(projects["Job Cost"].mean(), 2)
        average_bid_price = round(projects["Bid Price"].mean(), 2)
        avg_profit_percent = _safe_profit_percent(average_bid_price, average_job_cost)

        ridge_bid, ridge_cost, rf_bid, rf_cost = estimate_costs(
            models,
            sqft_value,
            description_value,
            structure_value,
        )

        ridge_profit_percent = _safe_profit_percent(ridge_bid, ridge_cost)
        rf_profit_percent = _safe_profit_percent(rf_bid, rf_cost)

        recent_history = projects.sort_values("Awarded Date").tail(5)
        history_lines = []
        for job_number, row in recent_history.iterrows():
            awarded_date = row["Awarded Date"]
            awarded_text = "N/A" if pd.isna(awarded_date) else awarded_date.strftime("%Y-%m-%d")
            history_lines.append(
                f"Job {job_number} | Date: {awarded_text} | SqFt: {int(row['SqFt']):,} | "
                f"Bid: ${row['Bid Price']:,.2f} | Cost: ${row['Job Cost']:,.2f} | "
                f"P/L: {row['Profit and Loss %']:.2f}%"
            )

        estimating_text = (
            f"{detailed_project_report}\n"
            "Historical Data Summary:\n"
            f"Historical SqFt Range: {lower_limit_value:,} to {upper_limit_value:,}\n"
            f"Matching Project Count: {len(projects)}\n"
            f"Average Bid Price: ${average_bid_price:,.2f}\n"
            f"Average Job Cost: ${average_job_cost:,.2f}\n"
            f"Average % Profit: {avg_profit_percent:.2f}%\n\n"
            "Machine Learning Estimates:\n"
            f"Ridge Bid Price: ${ridge_bid:,.2f}\n"
            f"Ridge Job Cost: ${ridge_cost:,.2f}\n"
            f"Ridge % Profit: {ridge_profit_percent:.2f}%\n\n"
            f"Random Forest Bid Price: ${rf_bid:,.2f}\n"
            f"Random Forest Job Cost: ${rf_cost:,.2f}\n"
            f"Random Forest % Profit: {rf_profit_percent:.2f}%\n\n"
            "Most Recent Matching Historical Projects:\n"
            + "\n".join(history_lines)
        )

        collected_data["Estimating Details"] = {
            "starting_text": "Project-Specific Estimating Details:",
            "user_input": estimating_text,
        }

        estimate_status_var.set(
            "Estimate saved to proposal data. Click Next to continue or run again to update values."
        )

        ################### Matplotlib Charts Below ######################
        # Updated bar chart to include 'Profit and Loss %' row.
        # most_recent = projects.tail()
        most_recent = projects.tail().copy() # Use .copy() to avoid SettingWithCopyWarning

        # Plot only the numeric bar-series columns.
        chart_columns = most_recent[['SqFt', 'Bid Price', 'Job Cost']]

        chart = chart_columns.plot(kind='bar')

        # Ensure 'Bid Price' and 'Job Cost' are cast to float before formatting.
        most_recent['Bid Price'] = most_recent['Bid Price'].astype(float)
        most_recent['Job Cost'] = most_recent['Job Cost'].astype(float)

        # Format 'Bid Price' and 'Job Cost' to be rounded to the nearest two digits
        most_recent['Formatted Bid Price'] = most_recent['Bid Price'].apply(lambda x: f"{x:,.2f}")
        most_recent['Formatted Job Cost'] = most_recent['Job Cost'].apply(lambda x: f"{x:,.2f}")

        # Add the table with all columns including 'Profit and Loss %'
        table_data = most_recent[['SqFt', 'Bid Price', 'Job Cost',  
                                    'Profit and Loss %']].T
        table_data.columns = most_recent.index

        table = plt.table(cellText=table_data.values, rowLabels=table_data.index, 
                            colLabels=table_data.columns, loc='bottom', 
                            cellLoc='center', rowLoc='center')
        table.auto_set_font_size(False)
        table.set_fontsize(10)
        table.scale(1, 1) # Adjust the scaling to match the chart size

        chart.set_title(f"Most Recent {description_value} Projects Around {sqft_value} sqft", 
                        fontsize=16)
        chart.set_xlabel("Job Numbers", fontsize=14)
        chart.set_ylabel("Amount", fontsize=14)
        chart.legend(loc='upper right')
        chart.grid(visible=True, axis="y")

        display_text = f"The average bid price is: ${average_bid_price:,.2f} \n"
        display_text += f"The average job cost is: ${average_job_cost:,.2f}"
        plt.text(.01, .97, display_text, ha='left', va='top', 
                    transform=chart.transAxes)

        chart.axes.get_xaxis().set_visible(False)
        # Adjust the layout to fit the table and chart
        plt.subplots_adjust(left=0.2, bottom=0.2) 

        # Add tooltips to display label and value.
        cursor = mplcursors.cursor(chart, hover=True)
        cursor.connect("add", lambda sel: sel.annotation.set_text(f'{sel.artist.get_label()}: {sel.target[1]:,.2f}'))

        # Set the window title
        plt.gcf().canvas.manager.set_window_title(f"{sqft_value}_{description_value}_{todays_date}")

        plt.show()
        
        # New line chart for trend over time
        projects.loc[:, 'Awarded Date'] = pd.to_datetime(projects['Awarded Date'])
        projects = projects.sort_values('Awarded Date')

        plt.figure()
        line1, = plt.plot(projects['Awarded Date'], projects['Bid Price'], 
                            label='Bid Price', marker='o')
        line2, = plt.plot(projects['Awarded Date'], projects['Job Cost'], 
                            label='Job Cost', marker='o')
        plt.title(f"Trend of {sqft_value} sqft {description_value} Projects Over Time")
        plt.xlabel("Awarded Date")
        plt.xticks(rotation=45)
        plt.ylabel("Amount")
        plt.legend()
        plt.grid(True)

        # Add tooltips to display label and value.
        cursor = mplcursors.cursor([line1, line2], hover=True)
        cursor.connect("add", lambda sel: sel.annotation.set_text(f'{sel.artist.get_label()}: {sel.target[1]:,.2f}'))

        plt.tight_layout()
        plt.gcf().canvas.manager.set_window_title(f"Trend of {sqft_value}_{description_value}_{todays_date}")
        plt.show()


    def save_estimating_page():
        if "Estimating Details" not in collected_data:
            collected_data["Estimating Details"] = {
                "starting_text": "Project-Specific Estimating Details:",
                "user_input": (
                    "Estimating inputs were not run before moving to the next section.\n"
                    f"Description: {description_var.get().strip() or 'Not provided'}\n"
                    f"Structure Type: {structure_var.get().strip() or 'Not provided'}\n"
                    f"Total SqFt: {sqft_var.get().strip() or 'Not provided'}\n"
                    f"Lower Limit: {lower_limit_var.get().strip() or 'Not provided'}\n"
                    f"Upper Limit: {upper_limit_var.get().strip() or 'Not provided'}\n"
                    f"Equipment Budget: {equipment_cost_var.get().strip() or 'Not provided'}\n"
                    f"Disposal Budget: {disposal_cost_var.get().strip() or 'Not provided'}\n"
                    f"Number of Days: {num_days_var.get().strip() or 'Not provided'}\n"
                    f"Number of Workers: {num_workers_var.get().strip() or 'Not provided'}"
                ),
            }
        show_page("Project-Specific Details")

    tk.Label(page, text="Description:",
              font=(FONT, FONT_SIZE)).grid(row=2, column=0, sticky="e", padx=10, pady=5)
    description_combo = ttk.Combobox(
        page,
        textvariable=description_var,
        state="readonly",
        values=["Interior Demolition", "Building Demo", "House Demo"],
        width=36,
    )
    description_combo.grid(row=2, column=1, sticky="w", padx=10, pady=5)
    description_combo.bind("<<ComboboxSelected>>", on_description_change)

    tk.Label(page, text="Structure Type:", 
             font=(FONT, FONT_SIZE)).grid(row=3, column=0, sticky="e", padx=10, pady=5)
    ttk.Combobox(
        page,
        textvariable=structure_var,
        state="readonly",
        values=["Concrete", "Wood", "Metal", "Brick or Block", "Other"],
        width=36,
    ).grid(row=3, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Total SqFt:", 
             font=(FONT, FONT_SIZE)).grid(row=4, column=0, sticky="e", padx=10, pady=5)
    tk.Entry(page, textvariable=sqft_var, 
             width=40).grid(row=4, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Lower SqFt Limit:", 
             font=(FONT, FONT_SIZE)).grid(row=5, column=0, sticky="e", padx=10, pady=5)
    tk.Entry(page, textvariable=lower_limit_var, 
             width=40).grid(row=5, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Upper SqFt Limit:", 
             font=(FONT, FONT_SIZE)).grid(row=6, column=0, sticky="e", padx=10, pady=5)
    tk.Entry(page, textvariable=upper_limit_var, 
             width=40).grid(row=6, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Total Equipment Budget:", 
             font=(FONT, FONT_SIZE)).grid(row=7, column=0, sticky="e", padx=10, pady=5)
    tk.Entry(page, textvariable=equipment_cost_var, 
             width=40).grid(row=7, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Total Disposal Budget:", 
             font=(FONT, FONT_SIZE)).grid(row=8, column=0, sticky="e", padx=10, pady=5)
    tk.Entry(page, textvariable=disposal_cost_var, 
             width=40).grid(row=8, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Total Number of Days:", 
             font=(FONT, FONT_SIZE)).grid(row=9, column=0, sticky="e", padx=10, pady=5)
    tk.Entry(page, textvariable=num_days_var, 
             width=40).grid(row=9, column=1, sticky="w", padx=10, pady=5)

    tk.Label(page, text="Total Number of Workers:", 
             font=(FONT, FONT_SIZE)).grid(row=10, column=0, sticky="e", padx=10, pady=5)
    tk.Entry(page, textvariable=num_workers_var, 
             width=40).grid(row=10, column=1, sticky="w", padx=10, pady=5)

    tk.Button(page, text="Run Estimate", 
              command=run_estimate_and_store).grid(row=11, column=1, padx=10, pady=10, sticky="w")

    tk.Label(
        page,
        textvariable=estimate_status_var,
        font=(FONT, FONT_SIZE),
        fg="darkgreen",
        wraplength=700,
        justify="left",
    ).grid(row=12, column=0, columnspan=2, padx=10, pady=10, sticky="w")

    tk.Button(page, text="Back", 
              command=lambda: show_page("Project Overview")).grid(row=13, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_estimating_page).grid(row=13, column=1, pady=20, sticky="e")


# Page 2: Project-Specific Details
def create_project_specific_details_page():
    """
    Create the "Project-Specific Details" page in the main GUI.

    This function creates the "Project-Specific Details" page, which allows 
    the user to input specific details about the project, such as the scope 
    of demolition, drawings, and specifications. The data entered on this 
    page is saved to the `collected_data` dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Project-Specific Details"] = page

    tk.Label(page, text="Project-Specific Details", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Add a text box for user input
    details_text = tk.Text(page, font=(FONT, FONT_SIZE), wrap="word", height=PAGE_HEIGHT, width=PAGE_WIDTH)
    details_text.grid(row=1, column=0, columnspan=2, padx=5, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "Clearly define demolition scope based on provided drawings, specifications, "
        "online property details, and user notes:\n\n\n"
        "Confirm specific drawing sheets and specification sections referenced:\n\n\n"
        "Identify scope gaps or unclear responsibilities:\n\n\n"
        "Permit responsibilities:\n\n\n"
        "Site access, fencing, erosion control:\n\n\n"
        "Temporary utilities and site controls:\n\n\n"
    )

    details_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    # Function to save the details and proceed to the next page
    def save_project_details():
        """
        Save the data entered on the "Project-Specific Details" page.

        This function collects the user input from the "Project-Specific Details" 
        page and saves it to the `collected_data` dictionary under the 
        "Project-Specific Details" key. After saving, it navigates to the 
        "Methodology and Approach" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = details_text.get("1.0", "end").strip()
        collected_data["Project-Specific Details"] = {
            "user_input": user_input,
        }
        show_page("Methodology and Approach")

    # Add navigation buttons
    tk.Button(page, text="Back", 
              command=lambda: show_page("Project Overview")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_project_details).grid(row=2, column=1, pady=20, sticky="e")


# Page 3: Methodology and Approach
def create_methodology_and_approach_page():
    """
    Create the "Methodology and Approach" page in the main GUI.

    This function creates the "Methodology and Approach" page, which allows 
    the user to describe the proposed demolition approach, safety measures, 
    and logistics. The data entered on this page is saved to the `collected_data` 
    dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Methodology and Approach"] = page

    tk.Label(page, text="Methodology and Approach",
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    methodology_text = tk.Text(page, font=(FONT, FONT_SIZE), wrap="word", height=PAGE_HEIGHT, width=PAGE_WIDTH)
    methodology_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "Describe the proposed demolition approach:\n\n\n"
        "Identify potential challenges and solutions:\n\n\n"
        "Outline safety measures and protocols:\n\n\n"
        "Discuss site access and logistics:\n\n\n"
        "Address any special considerations:\n\n\n"
        "Demolition means/methods (mechanical, hand demo, top-down approach):\n\n\n"
        "Worker safety, fall protection, dust control, and stability monitoring:\n\n\n"
        "Clarify structural sequencing or shoring responsibilities (Tinys or others?):\n\n\n"
    )
    methodology_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    def save_methodology():
        """
        Save the data entered on the "Methodology and Approach" page.

        This function collects the user input from the "Methodology and Approach" 
        page and saves it to the `collected_data` dictionary under the 
        "Methodology and Approach" key. After saving, it navigates to the 
        "Order of Operations" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = methodology_text.get("1.0", "end").strip()
        collected_data["Methodology and Approach"] = {
            "user_input": user_input,
        }
        show_page("Order of Operations")

    tk.Button(page, text="Back", 
              command=lambda: show_page("Project-Specific Details")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_methodology).grid(row=2, column=1, pady=20, sticky="e")


# Page 4: Order of Operations
def create_order_of_operations_page():
    """
    Create the "Order of Operations" page in the main GUI.

    This function creates the "Order of Operations" page, which allows the 
    user to outline the sequence of demolition activities, critical path 
    items, and any special considerations. The data entered on this page is 
    saved to the `collected_data` dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Order of Operations"] = page

    tk.Label(page, text="Order of Operations", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    order_text = tk.Text(page, font=(FONT, FONT_SIZE), wrap="word", height=PAGE_HEIGHT, width=PAGE_WIDTH)
    order_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "Outline the sequence of demolition activities:\n\n\n"
        "Identify critical path items:\n\n\n"
        "Discuss site access and logistics:\n\n\n"
        "Address any special considerations:\n\n\n"
        "Include any necessary permits or approvals:\n\n\n"
        "Clarify structural sequencing or shoring responsibilities (Company or others?):\n\n\n"
        "Identify special access, temporary utilities, traffic control, "
        "or structural stabilization needed prior to demolition start:\n\n\n"
    )
    order_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    def save_order_of_operations():
        """
        Save the data entered on the "Order of Operations" page.

        This function collects the user input from the "Order of Operations" page 
        and saves it to the `collected_data` dictionary under the 
        "Order of Operations" key. After saving, it navigates to the 
        "Alternates and Breakouts" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = order_text.get("1.0", "end").strip()
        collected_data["Order of Operations"] = {
            "user_input": user_input,
        }
        show_page("Alternates and Breakouts")

    tk.Button(page, text="Back", 
              command=lambda: show_page("Methodology and Approach")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_order_of_operations).grid(row=2, column=1, pady=20, sticky="e")


def create_alternates_and_breakouts_page():
    """
    Create the "Alternates and Breakouts" page in the main GUI.

    This function creates the "Alternates and Breakouts" page, which allows 
    the user to list any alternates, phases, or unit breakouts required for 
    the project. The data entered on this page is saved to the `collected_data` 
    dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Alternates and Breakouts"] = page

    tk.Label(page, text="Alternates and Breakouts", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Add a text box for user input
    alternates_text = Text(page, font=(FONT, FONT_SIZE), 
                           wrap="word", height=PAGE_HEIGHT, width=PAGE_WIDTH)
    alternates_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "List any alternates or breakouts:\n\n\n"
        "Identify any special considerations:\n\n\n"
        "Include any necessary permits or approvals:\n\n\n"
        "Clarify structural sequencing or shoring responsibilities (Company or others?):\n\n\n"
        "Identify special access, temporary utilities, traffic control, "
        "or structural stabilization needed prior to demolition start:\n\n\n"
        "Are there alternates, phases, or unit breakouts required per bid form/spec?:\n\n\n"
        "Identify dimensions, quantities, confirm exclusions/limitations per alternate:\n\n\n"
    )
    alternates_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    def save_alternates_and_breakouts():
        """
        Save the data entered on the "Alternates and Breakouts" page.

        This function collects the user input from the "Alternates and Breakouts" 
        page and saves it to the `collected_data` dictionary under the 
        "Alternates and Breakouts" key. After saving, it navigates to the 
        "Inclusions" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = alternates_text.get("1.0", "end").strip()
        collected_data["Alternates and Breakouts"] = {
            "user_input": user_input,
        }
        show_page("Inclusions")

    tk.Button(page, text="Back", 
              command=lambda: show_page("Order of Operations")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_alternates_and_breakouts).grid(row=2, column=1, pady=20, sticky="e")


def create_inclusions_page():
    """
    Create the "Inclusions" page in the main GUI.

    This function creates the "Inclusions" page, which allows the user to 
    list all items included in the scope of work. The data entered on this 
    page is saved to the `collected_data` dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Inclusions"] = page

    tk.Label(page, text="Inclusions", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Add a text box for user input
    inclusions_text = Text(page, font=(FONT, FONT_SIZE), 
                           wrap="word", height=PAGE_HEIGHT, width=PAGE_WIDTH)
    inclusions_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "1.) \n"
        "2.) \n"
        "3.) \n"
        "4.) \n"
        "5.) \n"
    )
    inclusions_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    def save_inclusions():
        """
        Save the data entered on the "Inclusions" page.

        This function collects the user input from the "Inclusions" page and saves 
        it to the `collected_data` dictionary under the "Inclusions" key. After 
        saving, it navigates to the "Exclusions" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = inclusions_text.get("1.0", "end").strip()
        collected_data["Inclusions"] = {
            "user_input": user_input,
        }
        show_page("Exclusions")
    
    tk.Button(page, text="Back", 
              command=lambda: show_page("Alternates and Breakouts")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_inclusions).grid(row=2, column=1, pady=20, sticky="e")


def create_exclusions_page():
    """
    Create the "Exclusions" page in the main GUI.

    This function creates the "Exclusions" page, which allows the user to 
    list all items excluded from the scope of work. The data entered on this 
    page is saved to the `collected_data` dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Exclusions"] = page

    tk.Label(page, text="Exclusions", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    exclusions_text = Text(page, font=(FONT, FONT_SIZE), 
                           wrap="word", height=PAGE_HEIGHT, width=PAGE_WIDTH)
    exclusions_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "Scanning, locating, or marking underground/embedded utilities:\n\n\n"
        "Utility cut, cap, isolation, or safe disconnect (M/E/P/FP/Data):\n\n\n"
        "Engineering, structural evaluation, shoring, or temp bracing:\n\n\n"
        "Hazardous materials removal (asbestos, lead, mold):\n\n\n"
        "Reinstallation of structural components (lintels, steel, cinch plates, etc.):\n\n\n"
        "Permitting costs (other than grading permit above):\n\n\n"
        "Patching/restoration of surfaces or finishes:\n\n\n"
        "Anything not explicitly listed in Inclusions:\n\n\n"
    )
    exclusions_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    def save_exclusions():
        """
        Save the data entered on the "Exclusions" page.

        This function collects the user input from the "Exclusions" page and saves 
        it to the `collected_data` dictionary under the "Exclusions" key. After 
        saving, it navigates to the "Equipment" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = exclusions_text.get("1.0", "end").strip()
        collected_data["Exclusions"] = {
            "user_input": user_input,
        }
        show_page("Equipment")

    tk.Button(page, text="Back", 
              command=lambda: show_page("Inclusions")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_exclusions).grid(row=2, column=1, pady=20, sticky="e")
    

equipment_book = EquipmentBook()


def create_equipment_page():
    """
    Create the "Equipment" page in the main GUI.

    This function creates the "Equipment" page, which allows the user to 
    list all equipment needed for the project and calculate the total 
    equipment cost. The data entered on this page is saved to the 
    `collected_data` dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    global equipment_text, total_equipment_cost

    total_equipment_cost = 0.0

    page = tk.Frame(root)
    pages["Equipment"] = page

    tk.Label(page, text="Equipment", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Add a text box for user input
    equipment_text = Text(page, font=(FONT, FONT_SIZE), 
                          wrap="word", height=PAGE_HEIGHT, width=PAGE_WIDTH)
    equipment_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "List all equipment needed for the project:\n\n"
    )
    equipment_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    def update_total_equipment_cost(new_total):
        global total_equipment_cost
        total_equipment_cost = new_total

    def save_equipment():
        """
        Save the data entered on the "Equipment" page.

        This function collects the user input from the "Equipment" page, including 
        the total equipment cost, and saves it to the `collected_data` dictionary 
        under the "Equipment" key. After saving, it navigates to the 
        "Terms and Conditions" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = equipment_text.get("1.0", "end").strip()
        collected_data["Equipment"] = {
            "user_input": user_input,
            "total_cost": f"Total Equipment Cost: ${total_equipment_cost:.2f}",
        }
        show_page("Terms and Conditions")

    tk.Button(page, text="Back", 
              command=lambda: show_page("Exclusions")).grid(row=2, column=0, pady=20, sticky="w")
    
    tk.Button(page, text="Select Equipment", 
              command=lambda: equipment_book.open_equipment_book(equipment_text, update_total_equipment_cost)).grid(row=2, column=1, padx=250, pady=20, sticky="e")

    tk.Button(page, text="Next", 
              command=save_equipment).grid(row=2, column=2, pady=20, sticky="e")


def create_terms_and_conditions_page():
    """
    Create the "Terms and Conditions" page in the main GUI.

    This function creates the "Terms and Conditions" page, which allows the 
    user to input the terms and conditions for the project. The data entered 
    on this page is saved to the `collected_data` dictionary when the user 
    clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Terms and Conditions"] = page

    tk.Label(page, text="Terms and Conditions", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Add a text box for user input
    terms_text = Text(page, font=(FONT, FONT_SIZE), wrap="word", 
                      height=PAGE_HEIGHT, width=PAGE_WIDTH)
    terms_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "Insert your terms and conditions here...\n"
    )
    terms_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    def save_terms_and_conditions():
        """
        Save the data entered on the "Terms and Conditions" page.

        This function collects the user input from the "Terms and Conditions" page 
        and saves it to the `collected_data` dictionary under the 
        "Terms and Conditions" key. After saving, it navigates to the 
        "Regulatory Compliance" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = terms_text.get("1.0", "end").strip()
        collected_data["Terms and Conditions"] = {
            "user_input": user_input,
        }
        show_page("Regulatory Compliance")

    tk.Button(page, text="Back", 
              command=lambda: show_page("Equipment")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_terms_and_conditions).grid(row=2, column=1, pady=20, sticky="e")


def create_regulatory_compliance_page():
    """
    Create the "Regulatory Compliance" page in the main GUI.

    This function creates the "Regulatory Compliance" page, which allows the 
    user to input details about compliance with OSHA, TOSHA, EPA, and other 
    regulations. The data entered on this page is saved to the `collected_data` 
    dictionary when the user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Regulatory Compliance"] = page

    tk.Label(page, text="Regulatory Compliance", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Add a text box for user input
    regulatory_text = Text(page, font=(FONT, FONT_SIZE), 
                           wrap="word", height=PAGE_HEIGHT, width=PAGE_WIDTH)
    regulatory_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "All applicable OSHA, TOSHA, and EPA regulations strictly followed.\n\n\n"
        "If no hazardous material survey:\n\n\n"
        "Recommend Owner perform a hazardous materials survey (asbestos, lead) "
        "for accurate pricing and compliance.\n\n\n"
        "For Residential Remodels specifically:\n"
        "Work will comply with EPA Renovation, Repair and Painting (RRP) Rule. "
        "Distribution of EPA pamphlet is recommended.\n\n\n"
    )

    regulatory_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    def save_regulatory_compliance():
        """
        Save the data entered on the "Regulatory Compliance" page.

        This function collects the user input from the "Regulatory Compliance" page 
        and saves it to the `collected_data` dictionary under the 
        "Regulatory Compliance" key. After saving, it navigates to the 
        "Final Checklist" page.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = regulatory_text.get("1.0", "end").strip()
        # Save the data to the collected_data dictionary
        collected_data["Regulatory Compliance"] = {
            "user_input": user_input,
        }
        show_page("Final Checklist")

    tk.Button(page, text="Back", 
              command=lambda: show_page("Terms and Conditions")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_regulatory_compliance).grid(row=2, column=1, pady=20, sticky="e")


def create_final_checklist_page():
    """
    Create the "Final Checklist" page in the main GUI.

    This function creates the "Final Checklist" page, which allows the user 
    to review and confirm key items before finalizing the proposal. The 
    checklist results are saved to the `collected_data` dictionary when the 
    user clicks "Next."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root, padx=75, pady=75)
    pages["Final Checklist"] = page

    tk.Label(page, text="Final Checklist", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Create BooleanVars for each checklist item
    checklist_vars = [
        BooleanVar(value=False),
        BooleanVar(value=False),
        BooleanVar(value=False),
        BooleanVar(value=False),
        BooleanVar(value=False),
        BooleanVar(value=False),
        BooleanVar(value=False),
    ]

    # Checklist items
    checklist_items = [
        "Confirm Bid Form presence and alternates/unit prices",
        "Drawings/specs/supporting docs uploaded",
        "Project type clearly selected",
        "Inclusions/exclusions aligned with drawings/specs",
        "Abatement status confirmed",
        "Special site items confirmed (fencing, erosion control, traffic plans)",
        "Methodology and sequencing reviewed (standard/custom)",
    ]

    # Add Checkbuttons for each checklist item
    for i, item in enumerate(checklist_items):
        Checkbutton(page, text=item, variable=checklist_vars[i], 
                    font=(FONT, FONT_SIZE), anchor="w", 
                    wraplength=450).grid(row=i + 1, column=0, columnspan=2, sticky="w", padx=10, pady=5)

    # Add a "Finish" button to save the checklist and proceed
    def save_final_checklist():
        """
        Save the data entered on the "Final Checklist" page.

        This function collects the checklist results from the "Final Checklist" 
        page and saves them to the `collected_data` dictionary under the 
        "Final Checklist" key. After saving, it navigates to the "Final Page."

        Parameters:
        None

        Returns:
        None
        """
        # Collect the checklist results
        formatted_results = "\n".join(
            f"{'[X]' if var.get() else '[ ]'} {item}" for var, item in zip(checklist_vars, checklist_items)
        )

        # Save the checklist results to the collected_data dictionary
        collected_data["Final Checklist"] = {
            "starting_text": "Final Checklist (Prior to Drafting Proposal):",
            "user_input": formatted_results,
        }
        show_page("Final Page")


    tk.Button(page, text="Back", 
              command=lambda: show_page("Regulatory Compliance")).grid(row=len(checklist_items) + 1, column=0, pady=20, sticky="w")
    tk.Button(page, text="Next", 
              command=save_final_checklist).grid(row=len(checklist_items) + 1, column=1, pady=20, sticky="e")


def create_final_page():
    """
    Create the "Final Page" in the main GUI.

    This function creates the "Final Page," which allows the user to add 
    final notes or comments and save the completed proposal to a Word 
    document. The data entered on this page is saved to the `collected_data` 
    dictionary, and the proposal is saved to a Word document when the user 
    clicks "Finish."

    Parameters:
    None

    Returns:
    None
    """
    page = tk.Frame(root)
    pages["Final Page"] = page

    tk.Label(page, text="Final Page", 
             font=(TITLE_FONT, TITLE_FONT_SIZE)).grid(row=0, column=0, columnspan=2, pady=10)

    # Add a text box for user input (optional, if you want to collect final notes)
    final_text = Text(page, font=(FONT, FONT_SIZE), wrap="word", 
                      height=PAGE_HEIGHT, width=PAGE_WIDTH)
    final_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

    # Insert the starting text into the text box
    starting_text = (
        "Congratulations! The Work Scope Bid Proposal is Complete.\n"
        "If you have any final notes or comments, please add them here...\n"
    )
    final_text.insert("1.0", starting_text)  # Insert text at the beginning of the text box

    # Add a "Finish" button to save details and close the window
    def save_final_page():
        """
        Save the data entered on the "Final Page" and generate the Word document.

        This function collects the user input from the "Final Page" and saves it 
        to the `collected_data` dictionary under the "Final Page" key. It then 
        generates a Word document containing all the collected data and saves it 
        to the specified file path.

        Parameters:
        None

        Returns:
        None
        """
        # Collect the user input and starting text
        user_input = final_text.get("1.0", "end").strip()
        # Save the data to the collected_data dictionary
        collected_data["Final Page"] = {
            "user_input": user_input,
        }

        # Retrieve the project name from the global variable
        project_name_value = project_name_var.get()

        file_path = os.path.join(downloads_folder, 
                                 f"{project_name_value}_Work_Scope_Bid_Proposal_({todays_date}).docx")
        save_to_word(collected_data, file_path)

    tk.Button(page, text="Back", 
              command=lambda: show_page("Final Checklist")).grid(row=2, column=0, pady=20, sticky="w")
    tk.Button(page, text="Finish", 
              command=save_final_page).grid(row=2, column=1, pady=20, sticky="e")


# Create all pages
create_project_overview_page()
create_project_estimating_page()
create_project_specific_details_page()
create_methodology_and_approach_page()
create_order_of_operations_page()
create_alternates_and_breakouts_page()
create_inclusions_page()
create_exclusions_page()
create_equipment_page()
create_terms_and_conditions_page()
create_regulatory_compliance_page()
create_final_checklist_page()
create_final_page()

# Show the first page
show_page("Project Overview")

# Start the main loop
root.mainloop()
